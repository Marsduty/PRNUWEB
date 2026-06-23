from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


TEMPLATE = Path(
    r"C:\Users\Marsduty\Desktop\个人资料\硕士项目\2026信网科创"
    r"\附件5：信息网络安全学院第十四届学生科研创新大赛学生科研立项申报书模板（B赛道）.docx"
)
OUTPUT = Path(r"G:\PRNUweb\docs\PRNU智能取证与大规模检索平台_学生科研立项申报书_含深度伪造检测模块版.docx")

FONT = "仿宋_GB2312"
TITLE = "基于PRNU设备指纹与深度伪造检测的智能取证与大规模检索平台"


def set_run_font(run, *, size: float = 12, bold: bool = False, font: str = FONT) -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def set_paragraph_text(paragraph, text: str, *, size: float = 12, bold: bool = False, align=None) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    if align is not None:
        paragraph.alignment = align


def clear_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        clear_paragraph(paragraph)
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._element
        p.getparent().remove(p)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_text(cell, text: str, *, align=WD_ALIGN_PARAGRAPH.CENTER, size: float = 12, bold: bool = False) -> None:
    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    for index, line in enumerate(text.split("\n")):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, size=size, bold=bold)


def add_cell_paragraph(cell, text: str, *, bold: bool = False, first_line: bool = True) -> None:
    paragraph = cell.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    if first_line:
        paragraph.paragraph_format.first_line_indent = Pt(24)
    run = paragraph.add_run(text)
    set_run_font(run, size=12, bold=bold)


def set_body_cell(cell, title: str, paragraphs: list[str]) -> None:
    clear_cell(cell)
    heading = cell.paragraphs[0]
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(0)
    run = heading.add_run(title)
    set_run_font(run, size=12, bold=True)
    for text in paragraphs:
        add_cell_paragraph(cell, text)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_all_table_fonts(doc: Document) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1
                    for run in paragraph.runs:
                        set_run_font(run, size=12)


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(TEMPLATE)

    # 封面信息。
    replacements = {
        12: "项目名称： 基于PRNU设备指纹与深度伪造检测的智能取证",
        13: "与大规模检索平台",
        14: "项目负责人：              待填写",
        15: "负责人专业：              待填写",
        16: "指导教师：                待填写",
        17: "申报日期：          2026年06月14日",
    }
    for index, text in replacements.items():
        if index < len(doc.paragraphs):
            set_paragraph_text(doc.paragraphs[index], text, size=14 if index in {12, 13} else 12, align=WD_ALIGN_PARAGRAPH.CENTER)

    table = doc.tables[0]
    set_cell_text(table.rows[0].cells[6], TITLE)
    set_cell_text(table.rows[1].cells[6], "2026年  7月  至   2027年  4月")

    set_cell_text(table.rows[2].cells[6], "待填写")
    set_cell_text(table.rows[2].cells[11], "待填写")
    set_cell_text(table.rows[2].cells[15], "待填写")
    set_cell_text(table.rows[2].cells[21], "待填写")
    set_cell_text(table.rows[3].cells[7], "待填写")
    set_cell_text(table.rows[3].cells[15], "待填写")
    set_cell_text(table.rows[3].cells[20], "待填写")
    set_cell_text(
        table.rows[4].cells[4],
        "负责人前期已完成PRNU基础链路代码整理与Web平台原型开发，具备Python图像处理、后端接口、任务队列、数据库建模和前端交互实现基础；已围绕设备指纹构建、PCE比对、任务管理、结果判定、Docker部署和数据集批量导入开展实践，并已形成深度伪造检测模块前期设计，能够承担项目总体设计、关键算法联调、系统集成与演示材料撰写工作。获奖情况待补充。",
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )

    team_rows = [
        ("1", "待填写", "待填写", "待填写", "待填写", "总体设计、后端接口、系统集成"),
        ("2", "待填写", "待填写", "待填写", "待填写", "PRNU算法联调与大规模指纹检索"),
        ("3", "待填写", "待填写", "待填写", "待填写", "深伪检测、质量门控与证据解释"),
        ("4", "待填写", "待填写", "待填写", "待填写", "前端界面、部署测试与文档材料"),
    ]
    for row_index, values in zip(range(6, 10), team_rows):
        _, name, gender, student_id, major, duty = values
        set_cell_text(table.rows[row_index].cells[3], name)
        set_cell_text(table.rows[row_index].cells[5], gender)
        set_cell_text(table.rows[row_index].cells[8], student_id)
        set_cell_text(table.rows[row_index].cells[10], major)
        set_cell_text(table.rows[row_index].cells[16], duty)

    set_cell_text(table.rows[11].cells[2], "待填写")
    set_cell_text(table.rows[11].cells[5], "待填写")
    set_cell_text(table.rows[11].cells[8], "待填写")
    set_cell_text(table.rows[11].cells[10], "待填写")
    set_cell_text(table.rows[11].cells[14], "数字取证、图像安全、人工智能安全")
    set_cell_text(table.rows[11].cells[16], "信息网络安全学院")
    set_cell_text(table.rows[11].cells[19], "待填写")
    set_cell_text(table.rows[14].cells[1], "第一项目指导教师长期关注网络空间安全、数字图像取证、人工智能安全或相关方向，能够在研究路线把关、数据合规使用、实验设计、成果凝练和项目验收材料撰写方面提供指导。具体业务特长与主要学术成果待补充。", align=WD_ALIGN_PARAGRAPH.LEFT)

    body = doc.tables[1]
    set_body_cell(
        body.rows[0].cells[0],
        "一、项目简介",
        [
            "本项目面向涉网图像取证、图像来源鉴别和数字内容真实性分析需求，建设集PRNU设备指纹取证、大规模指纹检索和深度伪造检测于一体的智能取证平台。平台一方面完成设备指纹录入、PRNU提取、PCE比对、结果判定和设备指纹管理，另一方面引入面向人脸图像的深度伪造检测能力，输出风险等级、可疑区域提示、结构化证据和报告说明，形成“设备来源一致性鉴别+内容真实性检测”的综合分析能力。",
            "项目拟采用Web化、任务队列化和模块化架构，将离线算法能力封装为可上传、可排队、可追踪、可复核、可部署的网站服务。PRNU模块重点解决图像是否倾向来自某一设备的问题；大规模检索模块重点解决设备指纹库扩展后的快速候选召回问题；深度伪造检测模块重点解决换脸、身份替换、属性编辑、表情驱动等静态人脸图像篡改风险提示问题。"
        ],
    )
    set_body_cell(
        body.rows[1].cells[0],
        "二、前期设计",
        [
            "项目总体定位为面向公安实战训练、涉网图像来源鉴别和数字证据辅助分析场景的智能取证平台。平台以PRNU设备指纹为主要技术基础，围绕“设备指纹录入—指纹提取—指纹比对—结果判定—证据复核”的业务闭环开展设计，目标是把原本离线运行的PRNU算法链路封装为可演示、可管理、可扩展的Web系统。现有基础已经形成PRNU指纹构建、指纹数据库比对、外来图像比对、结果判定、任务队列、设备指纹管理、Docker部署和多数据集批量导入等核心能力，后续在此基础上补充大规模指纹检索服务，并正式纳入深度伪造检测模块。",
            "需求层面，系统面向四类典型使用流程。第一类是设备指纹库建设流程，用户按设备录入品牌、型号、MAC地址、备注和参考图像，系统异步完成PRNU指纹构建并将指纹与设备档案绑定。第二类是指纹数据库比对流程，用户上传待检图像，系统提取待检图像指纹并在库内检索可能同源设备，给出PCE得分、阈值判断和候选设备信息。第三类是外来图像比对流程，用户上传图像A和图像B，系统提取两张图像的PRNU特征并进行单图PCE比对。第四类是深度伪造检测流程，用户上传静态人脸图像，系统对每张人脸输出风险分数、风险等级、可疑区域提示、结构化证据和结果说明。",
            "总体架构采用“前端工作台+后端API+异步任务队列+算法Worker+数据库与对象存储”的分层设计。数据接入层负责参考图像、待检图像、外来比对图像和深伪检测图像上传，并完成文件类型、尺寸、数量和任务参数校验；业务应用层提供图像导入、PRNU提取、指纹比对、深伪检测、结果判定、指纹管理和统计看板页面；后端服务层负责设备档案、任务记录、结果记录、文件路径和接口权限管理；算法处理层完成图像中心裁剪、噪声残差提取、PRNU聚合、PCE/NCC计算、候选召回、深伪风险检测、可疑区域定位和阈值判定；数据与索引层保存设备信息、任务信息、参考图像、指纹文件、比对结果、深伪检测结果和检索索引；综合报告层将PRNU结论、检索候选和深伪检测证据按统一格式展示。",
            "数据结构设计上，系统以设备、设备指纹、任务、任务结果、深伪检测结果和文件对象为核心实体。设备表保存品牌、型号、系统生成设备名称、MAC地址、备注和创建时间；设备指纹表保存指纹编号、所属设备、参考图像数量、指纹文件路径、构建状态和更新时间；任务表保存任务类型、任务名称、提交时间、执行状态、任务号和错误信息；比对结果表保存候选设备、PCE、NCC、阈值、结论文本和详细匹配信息；深伪结果表保存文件哈希、人脸数量、逐人脸风险分数、风险等级、质量评价、可疑区域、证据标签、模型版本和处理耗时；文件对象表或对象存储路径记录原始上传图像、裁剪后图像、热力图、参考掩膜和指纹文件。通过这些实体可以保证每一次建库、重构、比对和深伪检测都有可追溯记录。",
            "PRNU设备指纹构建流程采用设备级聚合思路。用户上传同一设备拍摄的多张参考图像后，系统首先进行格式校验和方向筛选，再对图像执行1024×1024中心裁剪，禁止使用resize改变原始噪声纹理统计特征。随后算法提取每张图像的噪声残差，进行归一化、去内容干扰和聚合平均，得到该设备的稳定PRNU指纹。构建完成后，系统把指纹文件、参考图像数量、设备信息和任务记录一并保存；当用户修改设备参考图像或设备信息时，系统触发设备指纹重构任务，保证指纹库内容与设备档案同步。",
            "指纹数据库比对流程采用阈值判定和候选解释相结合的方式。用户上传待检图像后，后端先执行同样的中心裁剪和PRNU提取流程，生成查询指纹；随后与数据库中的设备指纹进行PCE比对。若存在超过阈值60的候选，系统输出“倾向认定设备指纹：设备名称与待检图像同源”，并展示超过阈值的匹配设备、PCE、NCC、参考图像数量和设备详情；若没有候选超过阈值，则只展示最高得分候选和“库中未检索到匹配设备”的结论，避免把低分候选误读为有效同源证据。",
            "外来图像比对流程用于不依赖设备库的两图同源判断。系统分别对图像A和图像B进行中心裁剪和PRNU提取，再计算两者之间的PCE和NCC。若PCE超过阈值60，输出倾向认定图像A和图像B同源；若未超过阈值，则输出倾向认定图像A和图像B不同源或未达到同源判定阈值。该流程适合临时检材之间的快速比对，也可作为数据库比对结果的辅助复核手段。",
            "大规模指纹检索模块是后续重点扩展内容。随着设备指纹库规模增大，如果每次查询都对库内全部设备执行完整PCE计算，计算耗时会随库规模线性增长，难以满足交互式检索需求。因此设计采用“粗检索召回、精比对复核、阈值判定输出”的两阶段结构：第一阶段从PRNU指纹中提取较低维、可索引的描述子，用于快速召回Top-K候选；第二阶段只对候选集合执行完整PCE精确比对，从而在保持判定可靠性的同时降低整体计算量。",
            "粗检索描述子计划从频域统计、块级能量分布、方向纹理响应和降维后的PRNU向量中构建。索引组织上，优先考虑按图像尺寸、设备品牌、设备型号和指纹特征向量进行分层索引；算法选择上，可对比局部敏感哈希、近似最近邻检索、产品量化或向量数据库方案。召回阶段输出Top-K候选设备及粗检索相似度，精排阶段重新读取候选设备的完整PRNU指纹并计算PCE，最终按PCE得分和阈值给出结论。评测时重点关注召回率、Top-K命中率、平均检索耗时和精确复核后的误判情况。",
            "任务队列设计用于解决PRNU建库、重构、比对和深伪检测任务耗时较长的问题。前端提交任务后立即获得与任务列表一致的任务号，后端将任务写入数据库并推送到异步队列，由PRNU Worker、检索 Worker或深伪检测 Worker执行。前端任务列表按时间倒序实时刷新，并按子模块展示对应任务：图像导入显示设备指纹构建任务，指纹管理显示设备指纹重构任务，指纹比对显示数据库比对和外来图像比对任务，深伪检测显示人脸检测、风险评估和报告生成任务，结果判定显示最新综合结论。该设计既能避免网页请求阻塞，也便于用户追踪任务状态和历史结果。",
            "前端交互设计以主界面总览和流程步骤弹窗为核心。主界面展示当前数据库录入设备指纹数量、今日录入设备指纹数量、今日比对任务数量、今日比对命中数量、深伪检测任务数量和品牌分布图；流程区域包括图像导入、PRNU提取、指纹比对、深伪检测和结果判定等步骤，点击步骤后进入对应子模块。设备指纹管理独立为弹窗，支持按设备名称、品牌、型号和MAC地址搜索，支持查看、修改、删除和参考图像更新。比对和深伪结果以列表形式展示任务号、任务名、任务类型和结论摘要，点击条目后弹出详细信息，保证主界面简洁且结果可复核。",
            "部署与安全设计方面，系统采用Docker Compose组织前端、后端、数据库、缓存队列、对象存储和算法Worker，便于在本地、实验室服务器或云服务器上迁移。图像文件和指纹文件保存在对象存储或挂载卷中，数据库只保存结构化元数据和访问路径。后续部署到公网环境时，需要增加登录认证、上传大小限制、文件类型白名单、接口访问控制、数据库备份、对象存储容量监控和日志审计机制，避免测试数据泄露和任务资源被滥用。",
            "深度伪造检测模块定位为综合数字内容鉴定网站中的专项能力，首期聚焦JPG、PNG等静态人脸图像，主要覆盖换脸、身份替换、属性编辑和表情驱动等常见深伪形式。模块不直接替代司法鉴定结论，而是提供辅助筛查与复核结果；当图像质量过低、人脸过小、遮挡严重或证据不足时，系统返回“无法可靠判断”及原因，避免强制二分类。",
            "深伪检测流程划分为输入质量门控、场景识别与任务路由、多证据检测、可疑区域定位、证据约束解释、报告与服务接口六个环节。质量门控检查文件格式、图像解码、人脸数量、人脸尺寸、清晰度、遮挡、过曝欠曝和压缩程度；场景路由借鉴FakeShield中任务解耦和域标签思想，但不使用单一硬标签直接认定伪造类型，而是输出带置信度的候选路线，如疑似换脸、疑似属性编辑或低质量通用检测。",
            "多证据深伪检测子模块采用可替换的多分支结构。空间语义分支关注五官边缘、肤质、光照、反光和面部几何一致性；纹理与频率分支关注重采样、局部频谱和生成纹理异常；必要时接入PRNU模块提供的噪声一致性证据，但不替代PRNU的设备来源判断。各分支通过轻量融合层得到人脸级风险分数，并使用验证集进行阈值选择和概率校准。",
            "可疑区域定位与解释生成子模块用于增强结果可复核性。首期先使用检测器注意力图、类别激活图等方式形成热力图；具备稳定掩膜数据后，再引入分割模型生成参考掩膜。解释生成器只引用检测分数、质量信息、候选异常区域和预定义证据标签，不允许脱离结构化证据自由编造细节。报告文本区分观测事实和模型推断，低风险不表述为绝对真实，中高风险不表述为已经证实伪造。",
            "深伪模块对外输出统一结果对象，至少包括任务编号、文件哈希、图像质量评价、人脸数量、逐人脸风险分数、风险等级、可疑区域、结构化证据、结果不确定性、模型版本和处理时间。综合报告层并列展示PRNU、指纹检索和深伪检测证据，不简单平均不同模块分数，而是说明各自适用条件和限制。",
            "前期评测设计分为功能测试、算法测试和性能测试。功能测试关注图像导入、设备指纹构建、任务状态刷新、设备管理、结果弹窗和异常提示是否稳定；PRNU算法测试关注同源样本PCE分布、非同源样本PCE分布、阈值60下的命中情况和最高候选解释是否合理；深伪算法测试关注AUC、准确率、精确率、召回率、F1、等错误率、概率校准、定位IoU或区域命中率、解释一致性和压缩缩放等鲁棒性；性能测试关注单设备指纹构建耗时、单次数据库比对耗时、库规模增长后的检索耗时、深伪检测平均延迟和任务失败率。通过上述测试可以为后续优化算法参数、索引策略、模型版本和服务器配置提供依据。"
        ],
    )
    set_body_cell(
        body.rows[2].cells[0],
        "系统架构图（文字示意）",
        [
            "用户端：浏览器上传参考图像、待检图像、外来比对图像和待检测人脸图像。",
            "业务层：图像导入、PRNU提取、指纹比对、深伪检测、结果判定、设备指纹管理、大规模检索。",
            "服务层：前端工作台→后端API→任务队列→PRNU Worker/检索 Worker/深伪检测 Worker。",
            "PRNU算法层：中心裁剪与质量校验→噪声残差提取→设备PRNU聚合→粗检索召回→PCE精确复核→阈值判定。",
            "深伪算法层：质量门控→人脸检测→场景路由→多证据风险检测→可疑区域热力图→结构化证据解释。",
            "数据层：PostgreSQL保存设备、任务和结果；对象存储保存原始图像、指纹文件和热力图；检索索引保存指纹描述子和候选召回结构。"
        ],
    )
    set_body_cell(
        body.rows[3].cells[0],
        "三、项目实施方案及进度安排",
        [
            "实施方法上，项目采用工程实现与实验评测相结合的路线。工程部分基于现有Web平台继续迭代，采用前端工作台、后端API、异步任务队列、关系数据库、对象存储和算法Worker的组合方式；算法部分围绕PRNU提取、PCE匹配、候选召回、检索精排、深伪质量门控、多证据检测和可疑区域提示开展验证；评测部分使用公开数据集和自建测试样例检验指纹构建成功率、检索召回率、比对准确性、深伪检测鲁棒性和任务处理效率。",
            "研究范围包括五类核心问题：一是图像导入与设备指纹构建的规范化流程；二是数据库比对和外来图像比对的统一任务管理；三是面向规模化设备指纹库的快速候选检索与精确PCE复核；四是面向静态人脸图像的深度伪造风险检测、可疑区域提示和结构化证据输出；五是PRNU、指纹检索和深伪检测结果在综合报告中的并列展示与限制说明。",
            "资源支撑方面，项目已具备Python PRNU基础代码、Docker化Web工程、PostgreSQL/Redis/对象存储环境、前端可视化界面，以及FloreView、FODB、VISION等数据集的批量导入脚本和实验基础。深伪模块可选取公开人脸深伪数据集作为训练与测试来源，建立样本登记表，记录真实性、伪造方法、人物身份、来源数据集、人脸框、分辨率、压缩质量和增强方式。后续需要补充服务器部署环境、检索索引库或向量检索组件、深伪基线模型、测试样例集、演示视频录制环境和技术文档模板。",
            "团队分工拟定如下：负责人负责总体方案、后端接口、任务队列和系统集成；成员二负责PRNU算法联调、大规模指纹检索索引和召回排序实验；成员三负责深度伪造检测基线模型、质量门控、热力图和结构化证据接口；成员四负责前端页面、结果弹窗、服务器部署、测试报告、使用手册和演示视频。具体姓名、学号和专业信息待申报前补充。",
            "2026年7月至8月：完成需求分析、现有代码整理、数据库结构确认、深伪模块接口设计和项目申报材料完善，形成项目设计说明。2026年9月至10月：完善PRNU建库、比对、任务队列、结果判定和设备指纹管理功能，形成可运行平台原型。2026年11月至12月：设计并实现大规模指纹检索模块，完成特征描述子、索引构建、Top-K召回和PCE精排联调。2027年1月至2月：接入深度伪造检测最小可行版本，完成图像上传、人脸质量门控、基线检测器、逐人脸风险输出、热力图和模板化报告。2027年3月：完成系统联调、跨模块报告展示、数据集测试、异常处理和性能优化。2027年4月：整理演示视频、部署说明、使用手册、核心代码和技术文档，准备结项验收。"
        ],
    )
    set_body_cell(
        body.rows[4].cells[0],
        "四、特色与创新性",
        [
            "项目特色首先体现在将PRNU设备指纹取证流程做成完整Web平台，而不是停留在离线脚本层面。平台围绕图像导入、PRNU提取、指纹比对、结果判定和设备指纹管理组织功能模块，能够直观展示任务流、候选结果、设备信息和PCE/NCC指标，适合科研展示、教学训练和小规模实战验证。",
            "第二，项目强调面向规模化指纹库的检索能力。传统做法通常对查询图像与库内所有设备指纹逐一进行精确PCE比对，当设备规模上升时计算成本快速增加。本项目拟引入粗召回与精复核结合的检索架构，通过特征索引快速缩小候选范围，再对候选执行完整PCE计算，从而兼顾速度和判定可靠性。",
            "第三，项目保留取证可解释性。比对结论不只输出同源或不同源，还展示设备名称、品牌型号、MAC地址、指纹编号、参考图像数量、PCE和NCC等信息，便于复核证据链和定位误判风险。未达到阈值时只展示最高得分候选，避免把所有低分候选误解为有效匹配。",
            "第四，项目在工程上支持异步任务队列和可部署架构。PRNU建库与比对属于耗时任务，采用任务队列和Worker处理可以避免前端阻塞，也便于后续扩展GPU推理或独立检索服务。Docker化部署方式有利于在实验室服务器、演示环境和后续云服务器之间迁移。",
            "第五，项目纳入深度伪造检测模块，形成“设备来源一致性鉴别+内容真实性检测”的双线能力。PRNU解决图像是否倾向来自某设备的问题，深伪检测关注内容是否存在换脸、身份替换、属性编辑或生成篡改迹象，两者并列展示证据及适用条件，避免单一模型结论被过度解读。",
            "第六，深伪模块强调检测、定位和解释解耦。风险判断由可校准的检测模型承担，可疑区域由热力图或参考掩膜提供，说明文本由结构化证据约束生成，从设计上降低语言模型幻觉和误导性掩膜边界对取证结论的影响。"
        ],
    )
    set_body_cell(
        body.rows[5].cells[0],
        "五、最终成果与项目交付",
        [
            "最终成果为“基于PRNU设备指纹与深度伪造检测的智能取证与大规模检索平台”。系统应具备可运行、可演示、可部署的软件形态，至少包含图像导入、设备指纹构建、指纹数据库比对、外来图像比对、深度伪造检测、结果判定、设备指纹管理、任务队列、数据统计和大规模指纹检索原型能力。",
            "项目交付物包括：一是Web平台源代码和核心算法代码；二是Docker部署说明、环境变量说明和服务器部署建议；三是用户使用手册，覆盖设备录入、图像上传、指纹构建、任务查看、深伪检测、结果判定和设备管理流程；四是技术文档，说明PRNU处理流程、PCE阈值、指纹库结构、大规模检索设计、深伪检测流程、质量门控规则、风险等级、证据字段和模型版本记录；五是系统演示视频，展示从设备指纹录入到查询比对、深伪检测和综合结果查看的完整链路。",
            "验收时计划展示四类场景：设备指纹库构建演示、待检图像数据库检索演示、外来图像比对演示、静态人脸图像深度伪造检测演示。大规模检索模块展示索引构建、候选召回和精确复核过程。深度伪造检测模块展示上传检测、质量提示、逐人脸风险分数、风险等级、热力图或参考区域、结构化证据和报告说明。"
        ],
    )

    for row in body.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        row.height = None

    # 表格提示底色适当去除，统一字体。
    set_all_table_fonts(doc)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                set_cell_shading(cell, "FFFFFF")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
